"""
Post-processor for Pandoc-generated .docx files.

After Pandoc creates the Word document, this module opens it with
``python-docx`` and applies enhancements that Pandoc cannot do natively:

  • Table improvements  (auto-fit, header repeat, borders, alt-row shading)
  • Task-list checkboxes (Unicode → Wingdings symbol replacement)
  • Code block styling   (font, size, background, borders, keep-together)
"""

from __future__ import annotations

import re
from typing import Optional

from converter.config.settings import AppSettings

# python-docx is imported lazily in postprocess() so the module can be
# loaded even when python-docx is not installed (graceful degradation).


# ============================================================================
#  TABLE POST-PROCESSING
# ============================================================================

class TablePostProcessor:
    """Enhance every table in a python-docx Document."""

    def process(self, doc, settings: AppSettings) -> None:
        from docx.shared import Pt, Inches
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        for table in doc.tables:
            # ── Auto-fit width ────────────────────────────────────────
            if settings.table_autofit:
                self._set_autofit(table)

            # ── Repeat header row across pages ────────────────────────
            if settings.table_header_repeat and table.rows:
                self._set_header_repeat(table.rows[0])

            # ── Improved borders ──────────────────────────────────────
            self._apply_borders(table)

            # ── Alternating row shading ───────────────────────────────
            if settings.table_alternating_shading:
                self._apply_alternating_shading(table)

    @staticmethod
    def _set_autofit(table) -> None:
        """Set the table to auto-fit to window width."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        # Set table width to 100% (5000 = 100% in fiftieths of a percent)
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:w"), "5000")
        tblW.set(qn("w:type"), "pct")

        # Set layout to autofit
        tblLayout = tblPr.find(qn("w:tblLayout"))
        if tblLayout is None:
            tblLayout = OxmlElement("w:tblLayout")
            tblPr.append(tblLayout)
        tblLayout.set(qn("w:type"), "autofit")

    @staticmethod
    def _set_header_repeat(first_row) -> None:
        """Mark *first_row* to repeat at the top of each page."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        trPr = first_row._tr.get_or_add_trPr()
        tblHeader = trPr.find(qn("w:tblHeader"))
        if tblHeader is None:
            tblHeader = OxmlElement("w:tblHeader")
            trPr.append(tblHeader)

    @staticmethod
    def _apply_borders(table) -> None:
        """Apply consistent single-line borders to the entire table."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        # Remove existing borders
        existing = tblPr.find(qn("w:tblBorders"))
        if existing is not None:
            tblPr.remove(existing)

        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            elem = OxmlElement(f"w:{edge}")
            elem.set(qn("w:val"), "single")
            elem.set(qn("w:sz"), "4")          # 0.5pt (in eighth-points)
            elem.set(qn("w:space"), "0")
            elem.set(qn("w:color"), "BFBFBF")  # Light grey
            borders.append(elem)
        tblPr.append(borders)

    @staticmethod
    def _apply_alternating_shading(table) -> None:
        """Apply light grey background to even-indexed data rows."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        for idx, row in enumerate(table.rows):
            if idx == 0:
                # Header row — apply a slightly darker shade
                for cell in row.cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), "D9E2F3")  # Light blue header
                    tcPr.append(shd)
            elif idx % 2 == 0:
                # Even data rows — light grey
                for cell in row.cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), "F2F2F2")
                    tcPr.append(shd)


# ============================================================================
#  TASK LIST POST-PROCESSING
# ============================================================================

class TaskListPostProcessor:
    """
    Replace Unicode checkbox symbols with font-appropriate characters.

    Pandoc (with ``+task_lists``) or our preprocessor inserts ``☑`` / ``☐``.
    We ensure they render correctly in Word by applying the Segoe UI Symbol
    font (broadly available on Windows 7+).
    """

    _CHECKBOX_CHARS = {"☑", "☐", "✓", "✗"}

    def process(self, doc) -> None:
        from docx.shared import Pt

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                text = run.text
                if any(ch in text for ch in self._CHECKBOX_CHARS):
                    # Apply a font that renders these symbols well
                    run.font.name = "Segoe UI Symbol"
                    run.font.size = run.font.size or Pt(11)


# ============================================================================
#  CODE BLOCK POST-PROCESSING
# ============================================================================

class CodeBlockPostProcessor:
    """
    Override Pandoc's code block styling with user preferences.

    Identifies code paragraphs by their Word style name — Pandoc uses
    ``"Source Code"`` for fenced code blocks.
    """

    # Style names Pandoc may use for code blocks
    _CODE_STYLES = {"Source Code", "source code"}

    def process(self, doc, settings: AppSettings) -> None:
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        code_paragraphs = [
            p for p in doc.paragraphs
            if p.style and p.style.name in self._CODE_STYLES
        ]

        if not code_paragraphs:
            return

        resolved_font = settings.code_font

        for para in code_paragraphs:
            # ── Font family & size ────────────────────────────────────
            for run in para.runs:
                run.font.name = resolved_font
                run.font.size = Pt(settings.code_font_size)

            # ── Line spacing ──────────────────────────────────────────
            pf = para.paragraph_format
            pf.line_spacing = settings.code_line_spacing

            # ── Background shading ────────────────────────────────────
            if settings.code_bg_color:
                self._set_paragraph_shading(
                    para, settings.code_bg_color.lstrip("#")
                )

            # ── Border visibility ─────────────────────────────────────
            if settings.code_border_visible:
                self._set_paragraph_border(para)

        # ── Keep-together logic ───────────────────────────────────────
        self._apply_keep_together(code_paragraphs)

    @staticmethod
    def _set_paragraph_shading(para, hex_color: str) -> None:
        """Apply a background fill to a paragraph via XML."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        pPr = para._p.get_or_add_pPr()
        # Remove existing shading
        existing = pPr.find(qn("w:shd"))
        if existing is not None:
            pPr.remove(existing)

        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color.upper())
        pPr.append(shd)

    @staticmethod
    def _set_paragraph_border(para) -> None:
        """Add a thin border around the paragraph."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        pPr = para._p.get_or_add_pPr()
        existing = pPr.find(qn("w:pBdr"))
        if existing is not None:
            pPr.remove(existing)

        pBdr = OxmlElement("w:pBdr")
        for edge in ("top", "left", "bottom", "right"):
            elem = OxmlElement(f"w:{edge}")
            elem.set(qn("w:val"), "single")
            elem.set(qn("w:sz"), "4")
            elem.set(qn("w:space"), "4")
            elem.set(qn("w:color"), "CCCCCC")
            pBdr.append(elem)
        pPr.append(pBdr)

    @staticmethod
    def _apply_keep_together(code_paragraphs: list) -> None:
        """
        Apply ``keep_together`` for short code blocks (≤ 30 paragraphs).

        Pandoc emits one paragraph per line in a code block, all sharing
        the "Source Code" style.  We group consecutive code paragraphs
        and apply keep-together only when the block is reasonably short.
        """
        if not code_paragraphs:
            return

        # Group consecutive code paragraphs into blocks
        blocks: list[list] = []
        current_block: list = [code_paragraphs[0]]

        all_paragraphs = code_paragraphs[0]._p.getparent()
        para_elements = list(all_paragraphs)

        # Build index map for quick lookup
        element_indices = {id(p._p): i for i, p in enumerate(code_paragraphs)}

        for i in range(1, len(code_paragraphs)):
            prev_p = code_paragraphs[i - 1]
            curr_p = code_paragraphs[i]

            # Check if they're adjacent in the document
            try:
                prev_idx = list(all_paragraphs).index(prev_p._p)
                curr_idx = list(all_paragraphs).index(curr_p._p)
                if curr_idx == prev_idx + 1:
                    current_block.append(curr_p)
                else:
                    blocks.append(current_block)
                    current_block = [curr_p]
            except ValueError:
                blocks.append(current_block)
                current_block = [curr_p]

        blocks.append(current_block)

        # Apply keep-together to short blocks
        for block in blocks:
            if len(block) <= 30:
                for para in block:
                    para.paragraph_format.keep_together = True
                    para.paragraph_format.keep_with_next = True
                # Last paragraph shouldn't keep-with-next
                if block:
                    block[-1].paragraph_format.keep_with_next = False


# ============================================================================
#  ORCHESTRATION
# ============================================================================

def postprocess(doc_path: str, settings: AppSettings) -> None:
    """
    Open the Pandoc-generated ``.docx`` at *doc_path*, apply all
    post-processing enhancements, and save in-place.
    """
    try:
        from docx import Document
    except ImportError:
        # python-docx not installed — skip post-processing silently
        return

    doc = Document(doc_path)

    TablePostProcessor().process(doc, settings)
    TaskListPostProcessor().process(doc)
    CodeBlockPostProcessor().process(doc, settings)

    doc.save(doc_path)
