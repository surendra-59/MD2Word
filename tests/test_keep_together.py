import unittest
from docx import Document
from converter.core.postprocessor import CodeBlockPostProcessor

class TestKeepTogether(unittest.TestCase):
    def test_keep_together_short_block(self):
        doc = Document()
        # Add 10 paragraphs (short block)
        paragraphs = []
        for _ in range(10):
            p = doc.add_paragraph("code line")
            p.style = "Normal" # Mocking
            paragraphs.append(p)
            
        processor = CodeBlockPostProcessor()
        processor._apply_keep_together(paragraphs)
        
        # Check that keep_together is True for all
        for p in paragraphs:
            self.assertTrue(p.paragraph_format.keep_together)
            
        # Check keep_with_next is True for all but last
        for p in paragraphs[:-1]:
            self.assertTrue(p.paragraph_format.keep_with_next)
        self.assertFalse(paragraphs[-1].paragraph_format.keep_with_next)
        
    def test_keep_together_long_block(self):
        doc = Document()
        # Add 35 paragraphs (long block)
        paragraphs = []
        for _ in range(35):
            p = doc.add_paragraph("code line")
            p.style = "Normal" # Mocking
            paragraphs.append(p)
            
        processor = CodeBlockPostProcessor()
        processor._apply_keep_together(paragraphs)
        
        # Check that keep_together is NOT True
        for p in paragraphs:
            self.assertFalse(p.paragraph_format.keep_together)

if __name__ == '__main__':
    unittest.main()
